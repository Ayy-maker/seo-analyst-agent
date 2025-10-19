"""
Microsoft Word Document Parser
Supports .doc and .docx files
"""

from pathlib import Path
from typing import Dict, List
import re


class DOCXParser:
    """Parse Microsoft Word documents containing SEO data"""
    
    def parse(self, file_path: str) -> Dict:
        """Parse DOCX file and extract SEO data"""
        try:
            # Try to import docx library
            try:
                from docx import Document
            except ImportError:
                print("⚠️  python-docx not installed. Installing...")
                import subprocess
                subprocess.check_call(['pip', 'install', 'python-docx'])
                from docx import Document
            
            doc = Document(file_path)
            
            # Extract all text
            full_text = []
            tables_data = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
            
            # Extract tables (likely contains metrics)
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            # Try to identify report type and extract data
            text_content = '\n'.join(full_text)
            
            # Determine report type
            report_type = self._identify_report_type(text_content)
            
            # Extract structured data
            structured_data = self._extract_structured_data(tables_data, text_content)
            
            return {
                'source_file': Path(file_path).name,
                'type': report_type,
                'format': 'docx',
                'raw_text': text_content,
                'tables': tables_data,
                'structured_data': structured_data,
                'record_count': len(structured_data),
                'status': 'success'
            }
            
        except Exception as e:
            return {
                'source_file': Path(file_path).name,
                'error': f"Failed to parse DOCX: {str(e)}",
                'status': 'error'
            }
    
    def _identify_report_type(self, text: str) -> str:
        """Identify what type of SEO report this is"""
        text_lower = text.lower()
        
        if 'search console' in text_lower or 'google search' in text_lower:
            return 'search-console'
        elif 'analytics' in text_lower or 'traffic' in text_lower:
            return 'analytics'
        elif 'backlink' in text_lower or 'link' in text_lower:
            return 'backlinks'
        elif 'keyword' in text_lower or 'ranking' in text_lower:
            return 'keywords'
        elif 'technical' in text_lower or 'crawl' in text_lower:
            return 'technical'
        else:
            return 'general'
    
    def _extract_structured_data(self, tables: List, text: str) -> List[Dict]:
        """Extract structured data from tables and text"""
        data = []
        
        # Parse tables for metrics
        for table in tables:
            if len(table) < 2:
                continue
            
            # Assume first row is headers
            headers = table[0]
            
            for row in table[1:]:
                if len(row) == len(headers):
                    row_dict = dict(zip(headers, row))
                    data.append(row_dict)
        
        # If no tables, try to extract key metrics from text
        if not data:
            data = self._extract_metrics_from_text(text)
        
        return data
    
    def _extract_metrics_from_text(self, text: str) -> List[Dict]:
        """Extract metrics from plain text using patterns"""
        metrics = []
        
        # Look for common patterns like "Metric: Value"
        patterns = [
            r'(\w+(?:\s+\w+)*?):\s*([0-9,]+(?:\.\d+)?)',
            r'(\w+(?:\s+\w+)*?)\s*[-–]\s*([0-9,]+(?:\.\d+)?)',
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, text)
            for metric, value in matches:
                metrics.append({
                    'metric': metric.strip(),
                    'value': value.replace(',', '')
                })
        
        return metrics if metrics else [{'content': text[:500]}]
